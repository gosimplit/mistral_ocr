import os
import io
import re
import base64
import json
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import RGBColor, Inches
from werkzeug.datastructures import FileStorage

try:
    from PIL import Image
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

app = Flask(__name__)

@app.get("/health")
def health():
    return jsonify({"ok": True})

# ------------ Utilidades Markdown -> DOCX ----------------

HEADING_RE      = re.compile(r'^(#{1,6})\s+(.*)$')
UL_RE           = re.compile(r'^\s*[-*+]\s+(.*)$')
OL_RE           = re.compile(r'^\s*\d+\.\s+(.*)$')
TABLE_ROW_RE    = re.compile(r'^\s*\|(.+)\|\s*$')
IMG_LINE_RE     = re.compile(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$')
IMG_INLINE_RE   = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

def force_styles_black(doc: Document):
    target_styles = ["Normal", "List Paragraph", "List Bullet", "List Number"]
    target_styles += [f"Heading {i}" for i in range(1, 10)]
    for name in target_styles:
        try:
            style = doc.styles[name]
            if style and style.font:
                style.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

def add_paragraph(doc, text):
    if not text.strip():
        doc.add_paragraph("")
    else:
        doc.add_paragraph(text)

def flush_list(doc, buf, ordered):
    if not buf:
        return
    style = "List Number" if ordered else "List Bullet"
    for item in buf:
        doc.add_paragraph(item, style=style)
    buf.clear()

def is_align_row(row: str) -> bool:
    row = row.strip().strip("|").strip()
    cells = [c.strip() for c in row.split("|")]
    return all(re.fullmatch(r':?-{3,}:?', c) for c in cells)

def flush_table(doc, rows):
    if not rows:
        return
    filtered = [r for r in rows if not is_align_row(r)]
    if not filtered:
        return
    matrix = []
    for r in filtered:
        m = TABLE_ROW_RE.match(r)
        if not m:
            continue
        cells = [c.strip() for c in m.group(1).split("|")]
        matrix.append(cells)
    if not matrix:
        return
    cols = max(len(r) for r in matrix)
    table = doc.add_table(rows=len(matrix), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(matrix):
        for j in range(cols):
            table.cell(i, j).text = row[j] if j < len(row) else ""

def add_image_paragraph(doc: Document, img_bytes: bytes):
    section = doc.sections[-1]
    usable_width_emu = section.page_width - section.left_margin - section.right_margin
    EMUS_PER_INCH = 914400
    usable_width_in = float(usable_width_emu) / EMUS_PER_INCH

    stream = io.BytesIO(img_bytes)

    if PIL_AVAILABLE:
        try:
            with Image.open(io.BytesIO(img_bytes)) as im:
                width_px, height_px = im.size
                dpi_x = im.info.get("dpi", (96, 96))[0] or 96
                width_in = width_px / float(dpi_x)
                scale = 1.0
                if width_in > usable_width_in:
                    scale = usable_width_in / width_in
                new_width_in = width_in * scale
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(stream, width=Inches(new_width_in))
                return
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run()
    stream.seek(0)
    run.add_picture(stream, width=Inches(usable_width_in))

def handle_inline_images(doc: Document, text: str, images: dict):
    parts = []
    last_end = 0
    for m in IMG_INLINE_RE.finditer(text):
        if m.start() > last_end:
            parts.append(("text", text[last_end:m.start()]))
        fname = (m.group(2) or "").strip()
        parts.append(("img", fname))
        last_end = m.end()
    if last_end < len(text):
        parts.append(("text", text[last_end:]))

    if not parts:
        add_paragraph(doc, text)
        return

    for kind, payload in parts:
        if kind == "text":
            if payload.strip():
                add_paragraph(doc, payload.strip())
            else:
                doc.add_paragraph("")
        else:
            fname = payload
            blob = images.get(fname)
            if blob:
                add_image_paragraph(doc, blob)

def markdown_to_doc(md_text: str, images: dict, filename: str = "output.docx"):
    doc = Document()
    force_styles_black(doc)

    lines = md_text.splitlines()
    ul_buf, ol_buf, tbl_buf = [], [], []
    in_table = False
    para_buf = []

    def flush_para():
        if para_buf:
            text = " ".join(para_buf).strip()
            para_buf.clear()
            if IMG_INLINE_RE.search(text):
                handle_inline_images(doc, text, images)
            else:
                add_paragraph(doc, text)

    for raw in lines:
        line = raw.rstrip("\n")

        if TABLE_ROW_RE.match(line):
            in_table = True
            tbl_buf.append(line)
            continue
        else:
            if in_table:
                flush_para()
                flush_list(doc, ul_buf, ordered=False)
                flush_list(doc, ol_buf, ordered=True)
                flush_table(doc, tbl_buf)
                tbl_buf = []
                in_table = False

        m = HEADING_RE.match(line)
        if m:
            flush_para()
            flush_list(doc, ul_buf, ordered=False)
            flush_list(doc, ol_buf, ordered=True)
            level = len(m.group(1))
            text = m.group(2).strip()
            level = min(max(level, 1), 9)
            doc.add_heading(text, level=level)
            continue

        m_ul = UL_RE.match(line)
        m_ol = OL_RE.match(line)
        if m_ul:
            flush_para()
            flush_list(doc, ol_buf, ordered=True)
            ul_buf.append(m_ul.group(1).strip())
            continue
        if m_ol:
            flush_para()
            flush_list(doc, ul_buf, ordered=False)
            ol_buf.append(m_ol.group(1).strip())
            continue

        m_img = IMG_LINE_RE.match(line)
        if m_img:
            flush_para()
            flush_list(doc, ul_buf, ordered=False)
            flush_list(doc, ol_buf, ordered=True)
            fname = (m_img.group(2) or "").strip()
            blob = images.get(fname)
            if blob:
                add_image_paragraph(doc, blob)
            continue

        if not line.strip():
            flush_para()
            flush_list(doc, ul_buf, ordered=False)
            flush_list(doc, ol_buf, ordered=True)
            continue

        para_buf.append(line.strip())

    flush_para()
    flush_list(doc, ul_buf, ordered=False)
    flush_list(doc, ol_buf, ordered=True)
    if tbl_buf:
        flush_table(doc, tbl_buf)

    force_styles_black(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, filename

# -------------------- Endpoint principal --------------------

@app.post("/docx")
def make_docx():
    # --------- Caso JSON con imágenes en base64 ---------
    data = request.get_json(silent=True)
    if data and isinstance(data, dict) and ("markdown" in data or "text" in data):
        base_name = data.get("output_name") or data.get("filename") or "output"
        if not base_name.lower().endswith(".docx"):
            base_name += ".docx"

        images_map = {}
        if "images" in data and isinstance(data["images"], list):
            for img_obj in data["images"]:
                img_id = img_obj.get("id")
                img_b64 = img_obj.get("image_base64")
                if img_id and img_b64:
                    try:
                        images_map[img_id] = base64.b64decode(img_b64)
                    except Exception:
                        pass

        if data.get("markdown"):
            buf, fname = markdown_to_doc(data["markdown"], images=images_map, filename=base_name)
        else:
            doc = Document()
            force_styles_black(doc)
            add_paragraph(doc, data.get("text", ""))
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            fname = base_name

        return send_file(
            buf,
            as_attachment=True,
            download_name=fname,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # --------- Caso multipart con imágenes en string JSON ---------
    if request.form and ("markdown" in request.form or "text" in request.form):
        md_text = request.form.get("markdown", None)
        plain_text = request.form.get("text", None)
        base_name = request.form.get("output_name") or request.form.get("filename") or "output"
        if not base_name.lower().endswith(".docx"):
            base_name += ".docx"

        images_map = {}

        # Procesar campo "images" como string JSON
        if "images" in request.form:
            try:
                images_json = json.loads(request.form["images"])
                if isinstance(images_json, list):
                    for img_obj in images_json:
                        img_id = img_obj.get("id")
                        img_b64 = img_obj.get("image_base64")
                        if img_id and img_b64:
                            images_map[img_id] = base64.b64decode(img_b64)
            except Exception:
                pass

        # También permitir ficheros binarios (opcional)
        for f in request.files.getlist("file"):
            if isinstance(f, FileStorage) and f.filename:
                images_map[f.filename] = f.read()

        if md_text:
            buf, fname = markdown_to_doc(md_text, images_map, filename=base_name)
        else:
            doc = Document()
            force_styles_black(doc)
            add_paragraph(doc, plain_text or "")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            fname = base_name

        return send_file(
            buf,
            as_attachment=True,
            download_name=fname,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # --------- Si no es ni JSON ni multipart ---------
    return jsonify({
        "error": "Bad request: envía JSON (markdown + images base64) o multipart/form-data (markdown + campo images como string JSON)."
    }), 400

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
